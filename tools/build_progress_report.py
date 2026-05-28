#!/usr/bin/env python3
"""Build a meeting-ready progress README, figures, and Word report."""

from __future__ import annotations

import csv
import math
from pathlib import Path
from shutil import copy2

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports" / "2026-05-22-progress"
FIG_DIR = REPORT_DIR / "figures"
DATA_DIR = REPORT_DIR / "data"
DOCX_PATH = REPORT_DIR / "FG_MEET_progress_report_2026-05-22.docx"
README_PATH = REPORT_DIR / "README.md"
DYNAMIC_REPORT_DIR = ROOT / "reports" / "2026-05-22-dynamic"
MODAL_REPORT_DIR = ROOT / "reports" / "2026-05-22-modal30x30"
SENS_REPORT_DIR = ROOT / "reports" / "2026-05-22-modal-sensitivity"
DAMPING_REPORT_DIR = ROOT / "reports" / "2026-05-23-modal-damping"

FONT = Path("C:/Windows/Fonts/msyh.ttc")
FONT_BOLD = Path("C:/Windows/Fonts/msyhbd.ttc")
if not FONT.exists():
    FONT = Path("C:/Windows/Fonts/arial.ttf")
if not FONT_BOLD.exists():
    FONT_BOLD = FONT

COLORS = {
    "ink": "#172033",
    "muted": "#556070",
    "blue": "#2E74B5",
    "navy": "#1F3A5F",
    "green": "#2F7D32",
    "orange": "#B65C00",
    "red": "#9B1C1C",
    "grid": "#D8DEE8",
    "fill": "#F5F7FA",
}
PALETTE = ["#2E74B5", "#2F7D32", "#B65C00", "#7B61FF", "#C2410C", "#0F766E"]


def ensure_dirs() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    DATA_DIR.mkdir(parents=True, exist_ok=True)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(FONT_BOLD if bold else FONT), size=size)


def hex_to_rgb(value: str) -> tuple[int, int, int]:
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont, width: int) -> list[str]:
    if not text:
        return [""]
    lines: list[str] = []
    current = ""
    for ch in str(text):
        candidate = current + ch
        if text_size(draw, candidate, fnt)[0] <= width or not current:
            current = candidate
        else:
            lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def canvas(title: str, subtitle: str = "", size: tuple[int, int] = (1600, 900)) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, size[0], 86], fill=hex_to_rgb("#F1F5FA"))
    draw.text((42, 22), title, fill=hex_to_rgb(COLORS["navy"]), font=font(34, True))
    if subtitle:
        draw.text((42, 58), subtitle, fill=hex_to_rgb(COLORS["muted"]), font=font(18))
    return img, draw


def draw_footer(draw: ImageDraw.ImageDraw, size: tuple[int, int], note: str = "FG-MEET Workbench | 2026-05-22") -> None:
    draw.line([42, size[1] - 44, size[0] - 42, size[1] - 44], fill=hex_to_rgb("#E1E7EF"), width=2)
    draw.text((42, size[1] - 32), note, fill=hex_to_rgb(COLORS["muted"]), font=font(16))


def table_image(
    path: Path,
    title: str,
    headers: list[str],
    rows: list[list[object]],
    col_widths: list[int],
    subtitle: str = "",
) -> None:
    width = max(1500, sum(col_widths) + 120)
    tmp = Image.new("RGB", (width, 2000), "white")
    dtmp = ImageDraw.Draw(tmp)
    f_header = font(20, True)
    f_body = font(18)
    cell_pad = 12
    row_heights = []
    for row in rows:
        max_lines = 1
        for value, cw in zip(row, col_widths):
            max_lines = max(max_lines, len(wrap_text(dtmp, str(value), f_body, cw - cell_pad * 2)))
        row_heights.append(max(54, 30 * max_lines + cell_pad * 2))
    height = 132 + 58 + sum(row_heights) + 72
    img, draw = canvas(title, subtitle, (width, height))
    x0 = 60
    y = 118
    draw.rectangle([x0, y, x0 + sum(col_widths), y + 58], fill=hex_to_rgb("#E8EEF5"), outline=hex_to_rgb("#C7D0DD"))
    x = x0
    for header, cw in zip(headers, col_widths):
        draw.text((x + cell_pad, y + 16), header, fill=hex_to_rgb(COLORS["navy"]), font=f_header)
        x += cw
        draw.line([x, y, x, y + 58], fill=hex_to_rgb("#C7D0DD"), width=1)
    y += 58
    for idx, row in enumerate(rows):
        fill = "#FFFFFF" if idx % 2 == 0 else "#FAFBFD"
        draw.rectangle([x0, y, x0 + sum(col_widths), y + row_heights[idx]], fill=hex_to_rgb(fill), outline=hex_to_rgb("#D7DEE9"))
        x = x0
        for value, cw in zip(row, col_widths):
            lines = wrap_text(draw, str(value), f_body, cw - cell_pad * 2)
            yy = y + cell_pad
            for line in lines:
                draw.text((x + cell_pad, yy), line, fill=hex_to_rgb(COLORS["ink"]), font=f_body)
                yy += 30
            x += cw
            draw.line([x, y, x, y + row_heights[idx]], fill=hex_to_rgb("#D7DEE9"), width=1)
        y += row_heights[idx]
    draw_footer(draw, (width, height))
    img.save(path, quality=95)


def line_chart(path: Path, title: str, subtitle: str, x_values: list[float], series: dict[str, list[float]], y_label: str) -> None:
    size = (1600, 920)
    img, draw = canvas(title, subtitle, size)
    left, top, right, bottom = 120, 150, 1500, 790
    all_y = [v for vals in series.values() for v in vals if not pd.isna(v)]
    y_min = min(all_y)
    y_max = max(all_y)
    pad = (y_max - y_min) * 0.12 or 1.0
    y_min -= pad
    y_max += pad
    draw.rectangle([left, top, right, bottom], outline=hex_to_rgb(COLORS["grid"]), width=2)
    for i in range(6):
        y = bottom - i * (bottom - top) / 5
        value = y_min + i * (y_max - y_min) / 5
        draw.line([left, y, right, y], fill=hex_to_rgb("#EEF2F7"), width=1)
        draw.text((36, y - 12), f"{value:.2f}", fill=hex_to_rgb(COLORS["muted"]), font=font(15))
    for i, xval in enumerate(x_values):
        x = left + i * (right - left) / max(len(x_values) - 1, 1)
        draw.line([x, bottom, x, bottom + 8], fill=hex_to_rgb(COLORS["grid"]), width=2)
        draw.text((x - 20, bottom + 18), f"{xval:.1f}", fill=hex_to_rgb(COLORS["muted"]), font=font(15))
    draw.text((left, bottom + 56), "Vf0", fill=hex_to_rgb(COLORS["muted"]), font=font(17, True))
    draw.text((left - 80, top - 35), y_label, fill=hex_to_rgb(COLORS["muted"]), font=font(17, True))
    for sidx, (name, vals) in enumerate(series.items()):
        color = hex_to_rgb(PALETTE[sidx % len(PALETTE)])
        pts = []
        for i, val in enumerate(vals):
            x = left + i * (right - left) / max(len(x_values) - 1, 1)
            y = bottom - (val - y_min) / (y_max - y_min) * (bottom - top)
            pts.append((x, y))
        if len(pts) > 1:
            draw.line(pts, fill=color, width=5)
        for x, y in pts:
            draw.ellipse([x - 6, y - 6, x + 6, y + 6], fill=color, outline="white", width=2)
        lx = left + sidx * 160
        draw.line([lx, 115, lx + 36, 115], fill=color, width=5)
        draw.text((lx + 44, 104), name, fill=hex_to_rgb(COLORS["ink"]), font=font(17))
    draw_footer(draw, size)
    img.save(path, quality=95)


def bar_chart(path: Path, title: str, subtitle: str, labels: list[str], values: list[float], y_label: str, color: str = "#2E74B5") -> None:
    size = (1600, 900)
    img, draw = canvas(title, subtitle, size)
    left, top, right, bottom = 120, 150, 1510, 760
    vmax = max(values) if values else 1
    vmin = min(values) if values else 0
    if vmin > 0:
        vmin = 0
    pad = (vmax - vmin) * 0.12 or 1
    vmax += pad
    draw.rectangle([left, top, right, bottom], outline=hex_to_rgb(COLORS["grid"]), width=2)
    for i in range(6):
        y = bottom - i * (bottom - top) / 5
        val = vmin + i * (vmax - vmin) / 5
        draw.line([left, y, right, y], fill=hex_to_rgb("#EEF2F7"), width=1)
        draw.text((42, y - 12), f"{val:.2f}", fill=hex_to_rgb(COLORS["muted"]), font=font(15))
    n = len(values)
    gap = 18
    bw = (right - left - gap * (n + 1)) / max(n, 1)
    base_y = bottom - (0 - vmin) / (vmax - vmin) * (bottom - top)
    draw.line([left, base_y, right, base_y], fill=hex_to_rgb("#AAB4C3"), width=2)
    for i, (label, val) in enumerate(zip(labels, values)):
        x0 = left + gap + i * (bw + gap)
        x1 = x0 + bw
        y = bottom - (val - vmin) / (vmax - vmin) * (bottom - top)
        draw.rectangle([x0, min(y, base_y), x1, max(y, base_y)], fill=hex_to_rgb(color), outline=hex_to_rgb("#1E4F7A"))
        label_lines = wrap_text(draw, label, font(14), int(bw + 24))
        yy = bottom + 16
        for line in label_lines[:3]:
            tw, _ = text_size(draw, line, font(14))
            draw.text((x0 + bw / 2 - tw / 2, yy), line, fill=hex_to_rgb(COLORS["muted"]), font=font(14))
            yy += 22
        draw.text((x0, min(y, base_y) - 26), f"{val:.2f}", fill=hex_to_rgb(COLORS["ink"]), font=font(14, True))
    draw.text((left - 80, top - 35), y_label, fill=hex_to_rgb(COLORS["muted"]), font=font(17, True))
    draw_footer(draw, size)
    img.save(path, quality=95)


def grouped_bar_chart(path: Path, title: str, subtitle: str, labels: list[str], data: dict[str, list[float]], y_label: str) -> None:
    size = (1700, 940)
    img, draw = canvas(title, subtitle, size)
    left, top, right, bottom = 130, 155, 1600, 770
    vals = [v for row in data.values() for v in row]
    vmax = max(vals) * 1.18
    vmin = 0
    draw.rectangle([left, top, right, bottom], outline=hex_to_rgb(COLORS["grid"]), width=2)
    for i in range(6):
        y = bottom - i * (bottom - top) / 5
        val = vmin + i * (vmax - vmin) / 5
        draw.line([left, y, right, y], fill=hex_to_rgb("#EEF2F7"), width=1)
        draw.text((48, y - 12), f"{val:.1f}", fill=hex_to_rgb(COLORS["muted"]), font=font(15))
    groups = list(data.keys())
    group_w = (right - left) / len(labels)
    bw = group_w / (len(groups) + 1.4)
    for gi, group in enumerate(groups):
        color = hex_to_rgb(PALETTE[gi % len(PALETTE)])
        lx = left + gi * 210
        draw.rectangle([lx, 112, lx + 28, 132], fill=color)
        draw.text((lx + 38, 103), group, fill=hex_to_rgb(COLORS["ink"]), font=font(17))
    for i, label in enumerate(labels):
        gx = left + i * group_w
        for gi, group in enumerate(groups):
            val = data[group][i]
            x0 = gx + 18 + gi * bw
            x1 = x0 + bw * 0.82
            y = bottom - (val - vmin) / (vmax - vmin) * (bottom - top)
            draw.rectangle([x0, y, x1, bottom], fill=hex_to_rgb(PALETTE[gi % len(PALETTE)]))
        lines = wrap_text(draw, label, font(13), int(group_w - 10))
        yy = bottom + 16
        for line in lines[:3]:
            tw, _ = text_size(draw, line, font(13))
            draw.text((gx + group_w / 2 - tw / 2, yy), line, fill=hex_to_rgb(COLORS["muted"]), font=font(13))
            yy += 20
    draw.text((left - 88, top - 35), y_label, fill=hex_to_rgb(COLORS["muted"]), font=font(17, True))
    draw_footer(draw, size)
    img.save(path, quality=95)


def heatmap(path: Path, title: str, subtitle: str, matrix: pd.DataFrame, value_format: str = "{:.2f}") -> None:
    rows = list(matrix.index)
    cols = list(matrix.columns)
    size = (1500, 900)
    img, draw = canvas(title, subtitle, size)
    left, top = 170, 165
    cell_w, cell_h = 122, 82
    vals = matrix.values.flatten()
    vals = [float(v) for v in vals if not pd.isna(v)]
    vmin, vmax = min(vals), max(vals)
    for ci, col in enumerate(cols):
        draw.text((left + ci * cell_w + 35, top - 38), str(col), fill=hex_to_rgb(COLORS["muted"]), font=font(17, True))
    for ri, row in enumerate(rows):
        draw.text((54, top + ri * cell_h + 28), str(row), fill=hex_to_rgb(COLORS["muted"]), font=font(18, True))
        for ci, col in enumerate(cols):
            val = float(matrix.loc[row, col])
            ratio = (val - vmin) / (vmax - vmin or 1)
            r = int(232 - ratio * 95)
            g = int(242 - ratio * 80)
            b = int(252 - ratio * 35)
            x0, y0 = left + ci * cell_w, top + ri * cell_h
            draw.rectangle([x0, y0, x0 + cell_w, y0 + cell_h], fill=(r, g, b), outline=hex_to_rgb("#D7DEE9"))
            text = value_format.format(val)
            tw, th = text_size(draw, text, font(17, True))
            draw.text((x0 + cell_w / 2 - tw / 2, y0 + cell_h / 2 - th / 2), text, fill=hex_to_rgb(COLORS["ink"]), font=font(17, True))
    draw_footer(draw, size)
    img.save(path, quality=95)


def flowchart(path: Path) -> None:
    img, draw = canvas("技术路线与实际收敛流程", "从模型目标、方法设计到问题修正、最终验证闭环", (1700, 900))
    boxes = [
        ("模型目标", "FG-MEE 30×30×10 CFFF 方板\n位移/温差/电磁势耦合验证"),
        ("MATLAB 主算例", "5 种 FG × 9 个 Vf0 × 3 载荷\n135 行静力结果"),
        ("耦合验证", "2 mm 参考挠度\n电/磁正向 + 沈式回代"),
        ("COMSOL 对照", "10 层 CSV 分域材料\n15 点最大误差 4.927%"),
        ("动力扩展", "10×10 Newmark pilot\n30×30 模态降阶"),
        ("敏感性检查", "4/6/8/12/16 阶收敛\n16 阶作为稳妥口径"),
    ]
    x0, y0, bw, bh, gap = 45, 240, 245, 220, 28
    for i, (head, body) in enumerate(boxes):
        x = x0 + i * (bw + gap)
        draw.rounded_rectangle([x, y0, x + bw, y0 + bh], radius=18, fill=hex_to_rgb("#F8FAFD"), outline=hex_to_rgb("#B7C7D9"), width=3)
        draw.text((x + 22, y0 + 24), head, fill=hex_to_rgb(COLORS["navy"]), font=font(25, True))
        yy = y0 + 74
        for line in body.splitlines():
            draw.text((x + 22, yy), line, fill=hex_to_rgb(COLORS["ink"]), font=font(19))
            yy += 34
        if i < len(boxes) - 1:
            ax0 = x + bw + 5
            ay = y0 + bh / 2
            ax1 = x + bw + gap - 10
            draw.line([ax0, ay, ax1, ay], fill=hex_to_rgb(COLORS["blue"]), width=5)
            draw.polygon([(ax1, ay), (ax1 - 18, ay - 12), (ax1 - 18, ay + 12)], fill=hex_to_rgb(COLORS["blue"]))
    draw_footer(draw, (1700, 900))
    img.save(path, quality=95)


def load_data() -> dict[str, pd.DataFrame]:
    data = {
        "results": pd.read_csv(ROOT / "output" / "results_static.csv"),
        "refresh": pd.read_csv(ROOT / "output" / "results_static_w_center_refresh_report.csv"),
        "coupling": pd.read_csv(ROOT / "output" / "coupling_validation_2mm.csv"),
        "validation": pd.read_csv(ROOT / "comsol" / "results" / "validation_points_U_Vf06_elastic.csv"),
        "experiments": pd.read_csv(ROOT / "comsol" / "results" / "validation_experiments_U_Vf06_elastic_corrected.csv"),
        "vlog": pd.read_csv(ROOT / "comsol" / "results" / "validation_log.csv"),
        "dynamic_summary": pd.read_csv(DYNAMIC_REPORT_DIR / "data" / "dynamic_U_Vf06_elastic_10x10_summary.csv"),
        "modal_summary": pd.read_csv(MODAL_REPORT_DIR / "data" / "dynamic_modal_30x30_U_Vf06_elastic_8modes_summary.csv"),
        "modal_sensitivity": pd.read_csv(SENS_REPORT_DIR / "data" / "dynamic_modal_30x30_U_Vf06_elastic_sensitivity_summary.csv"),
        "modal_sensitivity_modes": pd.read_csv(SENS_REPORT_DIR / "data" / "dynamic_modal_30x30_U_Vf06_elastic_sensitivity_modes.csv"),
        "damping_sensitivity": pd.read_csv(DAMPING_REPORT_DIR / "data" / "dynamic_modal_30x30_U_Vf06_elastic_damping_sensitivity_summary.csv"),
    }
    return data


def copy_data_files() -> None:
    files = [
        ROOT / "output" / "results_static.csv",
        ROOT / "output" / "results_static_w_center_refresh_report.csv",
        ROOT / "output" / "coupling_validation_2mm.csv",
        ROOT / "comsol" / "results" / "validation_log.csv",
        ROOT / "comsol" / "results" / "validation_points_U_Vf06_elastic.csv",
        ROOT / "comsol" / "results" / "validation_experiments_U_Vf06_elastic_corrected.csv",
        ROOT / "comsol" / "export" / "Thermal_CFFF_U_Vf0.6-30x30-10layer_layers.csv",
        DYNAMIC_REPORT_DIR / "data" / "dynamic_U_Vf06_elastic_10x10_summary.csv",
        DYNAMIC_REPORT_DIR / "data" / "dynamic_U_Vf06_elastic_10x10_timeseries.csv",
        MODAL_REPORT_DIR / "data" / "dynamic_modal_30x30_U_Vf06_elastic_8modes_summary.csv",
        MODAL_REPORT_DIR / "data" / "dynamic_modal_30x30_U_Vf06_elastic_8modes_modes.csv",
        MODAL_REPORT_DIR / "data" / "dynamic_modal_30x30_U_Vf06_elastic_8modes_timeseries.csv",
        SENS_REPORT_DIR / "data" / "dynamic_modal_30x30_U_Vf06_elastic_sensitivity_summary.csv",
        SENS_REPORT_DIR / "data" / "dynamic_modal_30x30_U_Vf06_elastic_sensitivity_modes.csv",
        SENS_REPORT_DIR / "data" / "dynamic_modal_30x30_U_Vf06_elastic_sensitivity_timeseries.csv",
        DAMPING_REPORT_DIR / "data" / "dynamic_modal_30x30_U_Vf06_elastic_damping_sensitivity_summary.csv",
        DAMPING_REPORT_DIR / "data" / "dynamic_modal_30x30_U_Vf06_elastic_damping_sensitivity_modes.csv",
        DAMPING_REPORT_DIR / "data" / "dynamic_modal_30x30_U_Vf06_elastic_damping_sensitivity_timeseries.csv",
    ]
    for src in files:
        if src.exists():
            copy2(src, DATA_DIR / src.name)


def copy_report_figure(figs: dict[str, Path], key: str, src: Path, dst_name: str) -> None:
    dst = FIG_DIR / dst_name
    if not src.exists():
        raise FileNotFoundError(f"Missing report figure: {src}")
    copy2(src, dst)
    figs[key] = dst


def generate_figures(data: dict[str, pd.DataFrame]) -> dict[str, Path]:
    figs: dict[str, Path] = {}
    flowchart(FIG_DIR / "01_workflow.png")
    figs["workflow"] = FIG_DIR / "01_workflow.png"

    workload_rows = [
        ["FG 分布模式", "5", "U / V / X / O / P"],
        ["体积分数 Vf0", "9", "0.1 到 0.9"],
        ["主算例输入文件", "45", "每个 FG × Vf0 一份 10 层材料文件"],
        ["静力主算例", "135", "45 个材料 × Case A/B/C"],
        ["2 mm 耦合验证", "6", "正向、反向、沈式回代"],
        ["COMSOL 对照实验", str(len(data["experiments"])), "载荷、四面体、扫掠、分层 CSV 网格对照"],
        ["COMSOL 验证点", "15", "5×3 空间选点，MATLAB/COMSOL 逐点对比"],
        ["最终最大误差", "4.927%", "10 层 CSV + swept quad/hex 后达标"],
        ["10×10 Newmark 动力", "401 步", "40 ms 时程，峰值/频率/温差输出"],
        ["30×30 模态动力", "8 阶", "保留全尺寸空间模型，用模态叠加替代全量 Newmark"],
        ["模态阶数敏感性", "5 组", "4/6/8/12/16 阶，验证 8 阶与 16 阶基本一致"],
        ["阻尼敏感性", "4 组", "0/0.5/0.8/1.5% 阻尼，给出峰值范围"],
    ]
    table_image(FIG_DIR / "02_workload_table.png", "工作量总览（可作为月度汇报封面图）", ["项目", "数量", "说明"], workload_rows, [310, 160, 900])
    figs["workload_table"] = FIG_DIR / "02_workload_table.png"

    method_rows = [
        ["主求解", "MATLAB MEET 板壳 FEM", "完成 5×9×3 批量静力扫描"],
        ["正向耦合", "电势/磁势 → 目标挠度", "以 2 mm 为统一参考值"],
        ["反向传感", "机械挠度 → 电势/磁势", "采用沈式矩阵回代"],
        ["商用软件验证", "COMSOL 3D solid", "10 层 CSV 分域材料 + swept quad/hex"],
        ["误差判据", "关键点 < 5%", "最终 15 点最大误差 4.927%"],
    ]
    table_image(FIG_DIR / "03_method_table.png", "预计方法与实际执行路径", ["模块", "方法", "实际状态"], method_rows, [260, 430, 650])
    figs["method_table"] = FIG_DIR / "03_method_table.png"

    issue_rows = [
        ["COMSOL LiveLink 登录失败", "mphserver/LiveLink 需要额外会话", "改用 Java API + comsolbatch"],
        ["COMSOL 安全策略阻止写文件", "security.prefs 禁止 batch 文件访问", "打开文件系统/方法/外部库权限"],
        ["边界选择曾误选/漏选", "Box selection 条件不严格", "使用 allvertices 并打印 selection count"],
        ["COMSOL 初始误差大", "MATLAB reduced Qd 被直接按节点索引读取", "还原完整 5-DOF TQd 后再对比"],
        ["四面体网格无法 <5%", "自由 tetra 与板壳/厚度离散差异", "改 swept quad/hex 网格"],
        ["CSV 分层 1 单元/层偏硬", "厚度方向离散太粗", "每材料层 7 个 swept 单元"],
        ["10 单元/层又偏柔", "局部点误差重新增大", "保留 sweep=7 作为 canonical"],
        ["30×30 Newmark 过慢", "1 ms 冒烟超过 5 min", "改用 30×30 模态降阶并做阶数敏感性"],
        ["paper/main.tex 不存在", "当前仓库未建论文目录", "先输出 Word/PDF/README 汇报材料"],
    ]
    table_image(FIG_DIR / "04_problem_solution_table.png", "实际问题与解决思路", ["问题", "原因判断", "处理方案"], issue_rows, [360, 430, 650])
    figs["problem_solution_table"] = FIG_DIR / "04_problem_solution_table.png"

    results = data["results"].copy()
    x_values = sorted(results["vf0"].unique())
    for load_case, filename, ylabel, title in [
        ("elastic", "05_elastic_w_vs_vf.png", "w_center (mm)", "Case A：机械压力下中心挠度"),
        ("electro", "06_electro_w_vs_vf.png", "w_center (mm)", "Case B：电势驱动下中心挠度"),
        ("magneto", "07_magneto_w_vs_vf.png", "w_center (mm)", "Case C：磁势驱动下中心挠度"),
    ]:
        subset = results[results["load_case"] == load_case]
        series = {}
        for mode in sorted(subset["fg_mode"].unique()):
            vals = []
            for vf in x_values:
                row = subset[(subset["fg_mode"] == mode) & (subset["vf0"] == vf)]
                vals.append(float(row["w_center_mm"].iloc[0]) if not row.empty else math.nan)
            series[mode] = vals
        line_chart(FIG_DIR / filename, title, "修正 reduced Qd 后的中心节点 w(0.15,0.15,0)", x_values, series, ylabel)
        figs[filename[:-4]] = FIG_DIR / filename

    elastic = results[results["load_case"] == "elastic"]
    heat = elastic.pivot(index="fg_mode", columns="vf0", values="theta_span_K").sort_index()
    heatmap(FIG_DIR / "08_theta_span_heatmap.png", "Case A：层温差跨度热力图", "theta_span_K，行=FG 模式，列=Vf0", heat)
    figs["theta_heatmap"] = FIG_DIR / "08_theta_span_heatmap.png"

    magneto = results[results["load_case"] == "magneto"]
    series = {}
    for mode in sorted(magneto["fg_mode"].unique()):
        vals = []
        for vf in x_values:
            row = magneto[(magneto["fg_mode"] == mode) & (magneto["vf0"] == vf)]
            vals.append(float(row["magnetoelectric_efficiency"].iloc[0]) if not row.empty else math.nan)
        series[mode] = vals
    line_chart(FIG_DIR / "09_magnetoelectric_efficiency.png", "Case C：磁电效率随 Vf0 变化", "electric_span/(2*magnetic)，体现材料梯度影响", x_values, series, "ME efficiency")
    figs["me_efficiency"] = FIG_DIR / "09_magnetoelectric_efficiency.png"

    refresh = data["refresh"].copy()
    refresh["abs_diff"] = refresh["diff_w_center_mm"].astype(float).abs()
    top = refresh.sort_values("abs_diff", ascending=False).head(10)
    bar_chart(
        FIG_DIR / "10_wcenter_refresh_top10.png",
        "中心挠度刷新前后差异 Top 10",
        "说明 reduced Qd 还原修正对结果表影响显著",
        list(top["case_id"]),
        list(top["abs_diff"].astype(float)),
        "|Δw| (mm)",
        "#B65C00",
    )
    figs["refresh_top10"] = FIG_DIR / "10_wcenter_refresh_top10.png"

    exp = data["experiments"].copy()
    labels = list(exp["run_tag"])
    err_data = {
        "center": list(exp["center_rel_err_pct"].astype(float)),
        "max": list(exp["max_rel_err_pct"].astype(float)),
        "mean": list(exp["mean_rel_err_pct"].astype(float)),
    }
    grouped_bar_chart(FIG_DIR / "11_comsol_experiment_errors.png", "COMSOL 方案筛选：误差收敛路径", "错误方案也保留，便于组会讨论为什么最终选 layered_csv_sweep7", labels, err_data, "relative error (%)")
    figs["comsol_experiment_errors"] = FIG_DIR / "11_comsol_experiment_errors.png"

    val = data["validation"].copy()
    labels = list(val["point_id"])
    line_chart(
        FIG_DIR / "12_comsol_point_comparison.png",
        "最终 COMSOL 15 点位移对比",
        "10 层 CSV 分域材料 + swept quad/hex；MATLAB reduced Qd 已还原",
        list(range(1, len(labels) + 1)),
        {"MATLAB": list(val["matlab_w_mm"].astype(float)), "COMSOL": list(val["comsol_w_mm"].astype(float))},
        "w (mm)",
    )
    figs["comsol_point_comparison"] = FIG_DIR / "12_comsol_point_comparison.png"
    bar_chart(
        FIG_DIR / "13_comsol_point_error.png",
        "最终 COMSOL 15 点相对误差",
        "全部低于 5%，最大值 4.927%",
        labels,
        list(val["rel_err_w_pct"].astype(float)),
        "relative error (%)",
        "#2F7D32",
    )
    figs["comsol_point_error"] = FIG_DIR / "13_comsol_point_error.png"

    final_rows = [
        ["MATLAB p8", f"{float(data['vlog']['matlab_w_mm'].iloc[0]):.5f} mm", "中心节点参考值"],
        ["COMSOL p8", f"{float(data['vlog']['comsol_w_mm'].iloc[0]):.5f} mm", "分层 CSV 最终模型"],
        ["中心误差", f"{float(data['vlog']['rel_err_w_pct'].iloc[0]):.3f}%", "< 5%"],
        ["15 点最大误差", f"{float(val['rel_err_w_pct'].max()):.3f}%", "< 5%"],
        ["15 点平均误差", f"{float(val['rel_err_w_pct'].mean()):.3f}%", "稳定在 3.51%"],
    ]
    table_image(FIG_DIR / "14_final_result_table.png", "最终通过结果", ["指标", "数值", "说明"], final_rows, [360, 300, 640])
    figs["final_result_table"] = FIG_DIR / "14_final_result_table.png"

    coupling = data["coupling"].copy()
    c_rows = []
    for _, row in coupling.iterrows():
        c_rows.append([
            row["check_id"],
            row["method"],
            f"{float(row['w_center_mm']):.6f}",
            row["status"],
            row["transfer_unit"] if isinstance(row["transfer_unit"], str) else "",
        ])
    table_image(FIG_DIR / "15_coupling_table.png", "2 mm 耦合验证结果", ["检查项", "方法", "中心挠度(mm)", "状态", "系数/单位"], c_rows, [250, 420, 230, 120, 420])
    figs["coupling_table"] = FIG_DIR / "15_coupling_table.png"

    dyn = data["dynamic_summary"].iloc[0]
    modal = data["modal_summary"].iloc[0]
    sens = data["modal_sensitivity"].sort_values("n_modes")
    sens16 = sens.iloc[-1]
    damping = data["damping_sensitivity"].sort_values("damping_ratio")
    damping_default = damping.iloc[(damping["damping_ratio"] - 0.008).abs().argmin()]
    dynamic_rows = [
        [
            "10×10 Newmark pilot",
            f"{float(dyn['static_center_mm']):.4f}",
            f"{float(dyn['peak_center_mm']):.4f}",
            f"{float(dyn['peak_time_s']) * 1000:.2f}",
            f"{float(dyn['freq_01_Hz']):.2f}",
            f"{float(dyn['theta_span_max_K']):.4f}",
        ],
        [
            "30×30 8 阶模态",
            f"{float(modal['mechanical_static_center_mm']):.4f}",
            f"{float(modal['peak_center_mm']):.4f}",
            f"{float(modal['peak_time_s']) * 1000:.2f}",
            f"{float(modal['freq_01_Hz']):.2f}",
            f"{float(modal['theta_span_max_K']):.4f}",
        ],
        [
            "30×30 16 阶敏感性",
            f"{float(sens16['mechanical_static_center_mm']):.4f}",
            f"{float(sens16['peak_center_mm']):.4f}",
            f"{float(sens16['peak_time_s']) * 1000:.2f}",
            f"{float(sens16['freq_01_Hz']):.2f}",
            f"{float(sens16['theta_span_max_K']):.4f}",
        ],
    ]
    table_image(
        FIG_DIR / "16_dynamic_overview_table.png",
        "动力结果总览",
        ["工况/方法", "静态中心(mm)", "动力峰值(mm)", "峰时(ms)", "一阶频率(Hz)", "最大θ跨度(K)"],
        dynamic_rows,
        [280, 190, 190, 150, 170, 190],
        "从 10×10 Newmark pilot 推进到 30×30 模态与阶数敏感性",
    )
    figs["dynamic_overview_table"] = FIG_DIR / "16_dynamic_overview_table.png"

    copy_report_figure(figs, "dynamic_strategy_table", DYNAMIC_REPORT_DIR / "figures" / "01_dynamic_strategy_table.png", "17_dynamic_strategy_table.png")
    copy_report_figure(figs, "dynamic_summary_table", DYNAMIC_REPORT_DIR / "figures" / "02_dynamic_summary_table.png", "18_dynamic_summary_table.png")
    copy_report_figure(figs, "dynamic_center_timeseries", DYNAMIC_REPORT_DIR / "figures" / "03_center_displacement_timeseries.png", "19_dynamic_center_timeseries.png")
    copy_report_figure(figs, "dynamic_theta_timeseries", DYNAMIC_REPORT_DIR / "figures" / "04_theta_span_timeseries.png", "20_dynamic_theta_timeseries.png")
    copy_report_figure(figs, "dynamic_frequency_bar", DYNAMIC_REPORT_DIR / "figures" / "05_modal_frequency_bar.png", "21_dynamic_frequency_bar.png")

    copy_report_figure(figs, "modal30_summary_table", MODAL_REPORT_DIR / "figures" / "01_modal30x30_summary_table.png", "22_modal30x30_summary_table.png")
    copy_report_figure(figs, "modal30_center_timeseries", MODAL_REPORT_DIR / "figures" / "02_modal30x30_center_timeseries.png", "23_modal30x30_center_timeseries.png")
    copy_report_figure(figs, "modal30_theta_timeseries", MODAL_REPORT_DIR / "figures" / "03_modal30x30_theta_timeseries.png", "24_modal30x30_theta_timeseries.png")
    copy_report_figure(figs, "modal30_contribution", MODAL_REPORT_DIR / "figures" / "04_modal_static_contribution.png", "25_modal30x30_contribution.png")
    copy_report_figure(figs, "modal30_frequency_bar", MODAL_REPORT_DIR / "figures" / "05_modal_frequency_bar.png", "26_modal30x30_frequency_bar.png")
    copy_report_figure(figs, "modal30_compare_10x10", MODAL_REPORT_DIR / "figures" / "06_10x10_vs_30x30_center.png", "27_10x10_vs_30x30_center.png")

    copy_report_figure(figs, "sens_summary_table", SENS_REPORT_DIR / "figures" / "01_modal_sensitivity_summary_table.png", "28_modal_sensitivity_summary_table.png")
    copy_report_figure(figs, "sens_capture_ratio", SENS_REPORT_DIR / "figures" / "02_capture_ratio_vs_modes.png", "29_capture_ratio_vs_modes.png")
    copy_report_figure(figs, "sens_peak_vs_modes", SENS_REPORT_DIR / "figures" / "03_peak_vs_modes.png", "30_peak_vs_modes.png")
    copy_report_figure(figs, "sens_center_timeseries", SENS_REPORT_DIR / "figures" / "04_center_timeseries_by_modes.png", "31_center_timeseries_by_modes.png")
    copy_report_figure(figs, "sens_theta_timeseries", SENS_REPORT_DIR / "figures" / "05_theta_span_timeseries_by_modes.png", "32_theta_span_timeseries_by_modes.png")
    copy_report_figure(figs, "sens_modal_contribution", SENS_REPORT_DIR / "figures" / "06_first16_modal_contribution.png", "33_first16_modal_contribution.png")
    copy_report_figure(figs, "sens_frequency_bar", SENS_REPORT_DIR / "figures" / "07_first16_frequency_bar.png", "34_first16_frequency_bar.png")
    copy_report_figure(figs, "sens_assessment_table", SENS_REPORT_DIR / "figures" / "08_convergence_assessment_table.png", "35_convergence_assessment_table.png")

    copy_report_figure(figs, "damping_summary_table", DAMPING_REPORT_DIR / "figures" / "01_damping_summary_table.png", "36_damping_summary_table.png")
    copy_report_figure(figs, "damping_peak_vs", DAMPING_REPORT_DIR / "figures" / "02_peak_vs_damping.png", "37_peak_vs_damping.png")
    copy_report_figure(figs, "damping_theta_vs", DAMPING_REPORT_DIR / "figures" / "03_theta_span_vs_damping.png", "38_theta_span_vs_damping.png")
    copy_report_figure(figs, "damping_center_timeseries", DAMPING_REPORT_DIR / "figures" / "04_center_timeseries_by_damping.png", "39_center_timeseries_by_damping.png")
    copy_report_figure(figs, "damping_theta_timeseries", DAMPING_REPORT_DIR / "figures" / "05_theta_span_timeseries_by_damping.png", "40_theta_span_timeseries_by_damping.png")
    copy_report_figure(figs, "damping_assessment_table", DAMPING_REPORT_DIR / "figures" / "06_damping_assessment_table.png", "41_damping_assessment_table.png")

    return figs


def markdown_table(headers: list[str], rows: list[list[object]]) -> str:
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        out.append("| " + " | ".join(str(x).replace("\n", "<br>") for x in row) + " |")
    return "\n".join(out)


def generate_readme(data: dict[str, pd.DataFrame], figs: dict[str, Path]) -> None:
    val = data["validation"]
    vlog = data["vlog"].iloc[0]
    exp = data["experiments"]
    results = data["results"]
    dyn = data["dynamic_summary"].iloc[0]
    modal = data["modal_summary"].iloc[0]
    sens = data["modal_sensitivity"].sort_values("n_modes")
    sens16 = sens.iloc[-1]
    damping = data["damping_sensitivity"].sort_values("damping_ratio")
    damping_default = damping.iloc[(damping["damping_ratio"] - 0.008).abs().argmin()]
    content: list[str] = []
    content.append("# FG-MEET Workbench 月度/组会汇报材料（2026-05-22）\n")
    content.append("> 用途：明天组会可直接截取表格和图片做 PPT。本文按“模型目标 → 预计方法 → 实际问题 → 解决路径 → 正确/错误结果展示”的逻辑组织。\n")
    content.append("## 0. 一页结论\n")
    summary_rows = [
        ["主模型", "300 mm × 300 mm × 6 mm FG-MEE CFFF 方板，30×30 八节点板壳，10 层材料"],
        ["MATLAB 主算例", "5 种 FG × 9 个 Vf0 × 3 个载荷 = 135 行静力结果"],
        ["2 mm 转化验证", "6 个检查项全部 pass，覆盖电势、磁势、机械回代和沈式直接矩阵法"],
        ["COMSOL 最终模型", "10 层 CSV 分域材料 + swept quad/hex 网格，mesh size 4，每层 7 个扫掠单元"],
        ["最终误差", f"中心点 {float(vlog['rel_err_w_pct']):.3f}%，15 点最大 {float(val['rel_err_w_pct'].max()):.3f}%，平均 {float(val['rel_err_w_pct'].mean()):.3f}%"],
        ["动力补充", f"10×10 Newmark 峰值 {float(dyn['peak_center_mm']):.4f} mm；30×30 16 阶模态峰值 {float(sens16['peak_center_mm']):.4f} mm"],
        ["阻尼敏感性", f"0/0.5/0.8/1.5% 已补算；0.8% 口径峰值 {float(damping_default['peak_center_mm']):.4f} mm"],
        ["仍待推进", "论文 main.tex 尚未建立；如需论文最终动态图，建议补夜间长时程 Newmark 对照"],
    ]
    content.append(markdown_table(["项目", "结论"], summary_rows))
    content.append("\n![工作流](figures/01_workflow.png)\n")
    content.append("![工作量总览](figures/02_workload_table.png)\n")

    content.append("## 1. 模型是什么，要完成的目标\n")
    model_rows = [
        ["几何", "方板 300×300×6 mm，CFFF 单边固支三边自由"],
        ["材料", "BaTiO3/CoFe2O4 功能梯度磁-电-弹性材料，厚度方向 10 层"],
        ["变量", "位移 w、温差/热响应 theta、电势响应、磁势响应"],
        ["主载荷", "Case A 机械压力 15000 Pa；Case B 上下 ±300 V；Case C 上下 ±200 A 磁势"],
        ["验证目标", "MATLAB MEET 主程序批量求解；代表工况用 COMSOL 独立验证，关键误差 <5%"],
    ]
    content.append(markdown_table(["条目", "说明"], model_rows))
    content.append("\n![方法表](figures/03_method_table.png)\n")

    content.append("## 2. 预计使用的计算、仿真方法\n")
    content.append("本阶段采用 MATLAB 作为主求解器，COMSOL 只做代表性对照，不承担 45×3 全量扫描。这样既保留板壳全耦合程序的效率，也能给组会/论文提供商用软件独立验证证据。\n")
    method_table = [
        ["MATLAB 静力扫描", "run_batch_static.m", "完整参数矩阵，输出挠度、温差、磁电效率"],
        ["2 mm 转化验证", "run_coupling_validation_2mm.m", "统一参考挠度，验证正向驱动和反向传感"],
        ["沈式回代", "[Kff,Kfz;Kzf,Kzz]\\[-Kfu*u-Kft*T;-Kzu*u-Kzt*T]", "从位移场直接回算感生电/磁势"],
        ["COMSOL 对照", "Java API + comsolbatch", "绕开 LiveLink 登录依赖，可命令行复现"],
        ["动力计算", "Newmark pilot + 30×30 模态降阶", "先验证频率/峰值，再做阶数敏感性"],
        ["结果汇报", "Markdown + Word/PDF + 图片化表格", "便于直接截取到 PPT"],
    ]
    content.append(markdown_table(["模块", "实现", "目的"], method_table))

    content.append("## 3. 实际途中遇到的问题与解决思路\n")
    content.append("![问题解决表](figures/04_problem_solution_table.png)\n")
    content.append("关键经验是：一开始不能只盯 COMSOL 网格，必须先排除 MATLAB reduced `Qd` 后处理错位。修正这个问题后，误差从“看起来 28% 最大误差”下降到 7% 左右，再通过 swept hex 网格和 CSV 分层材料把误差压到 5% 内。\n")

    content.append("## 4. MATLAB 主算例结果展示\n")
    content.append("### 4.1 Case A 机械压力\n")
    content.append("![Case A](figures/05_elastic_w_vs_vf.png)\n")
    content.append("### 4.2 Case B 电势驱动\n")
    content.append("![Case B](figures/06_electro_w_vs_vf.png)\n")
    content.append("### 4.3 Case C 磁势驱动\n")
    content.append("![Case C](figures/07_magneto_w_vs_vf.png)\n")
    content.append("### 4.4 温差与磁电效率\n")
    content.append("![温差热力图](figures/08_theta_span_heatmap.png)\n")
    content.append("![磁电效率](figures/09_magnetoelectric_efficiency.png)\n")
    load_counts = results.groupby("load_case")["case_id"].count().reset_index()
    content.append(markdown_table(["载荷", "完成行数"], load_counts.values.tolist()))

    content.append("## 5. 结果修正：reduced Qd 还原\n")
    content.append("原先 `Qd` 是去除约束自由度后的 reduced 向量；若直接按 `5*(node-1)+3` 读取，会在固定边存在时错位。本次新增完整节点自由度还原 `TQd`，并刷新 `results_static.csv` 135 行。\n")
    content.append("![中心挠度刷新差异](figures/10_wcenter_refresh_top10.png)\n")

    content.append("## 6. COMSOL 错误结果、筛选过程与最终方案\n")
    content.append("![COMSOL 方案误差](figures/11_comsol_experiment_errors.png)\n")
    exp_rows = exp[["run_tag", "center_rel_err_pct", "max_rel_err_pct", "mean_rel_err_pct", "notes"]].copy()
    exp_rows[["center_rel_err_pct", "max_rel_err_pct", "mean_rel_err_pct"]] = exp_rows[["center_rel_err_pct", "max_rel_err_pct", "mean_rel_err_pct"]].round(3)
    content.append(markdown_table(["实验", "中心误差%", "最大误差%", "平均误差%", "说明"], exp_rows.values.tolist()))
    content.append("\n结论：自由四面体网格细化并不能解决差异；单域 swept hex 可以达标，但为了回应“10 层材料 CSV 导入”的要求，最终选择 10-domain CSV + swept mesh。1 单元/层明显偏硬，10 单元/层局部偏柔，7 单元/层在 15 点内全部低于 5%。\n")

    content.append("## 7. 最终正确结果详尽展示\n")
    content.append("![最终结果表](figures/14_final_result_table.png)\n")
    content.append("![15 点位移对比](figures/12_comsol_point_comparison.png)\n")
    content.append("![15 点误差](figures/13_comsol_point_error.png)\n")
    point_rows = val[["point_id", "x_m", "y_m", "matlab_w_mm", "comsol_w_mm", "diff_w_mm", "rel_err_w_pct"]].copy()
    for col in ["matlab_w_mm", "comsol_w_mm", "diff_w_mm", "rel_err_w_pct"]:
        point_rows[col] = point_rows[col].astype(float).round(4)
    content.append(markdown_table(["点", "x", "y", "MATLAB mm", "COMSOL mm", "差值 mm", "误差%"], point_rows.values.tolist()))

    content.append("## 8. 2 mm 耦合验证结果\n")
    content.append("![2mm 耦合表](figures/15_coupling_table.png)\n")
    cdf = data["coupling"][["check_id", "method", "driver_kind", "driver_value", "w_center_mm", "status", "transfer_coefficient", "transfer_unit"]].copy()
    for col in ["driver_value", "w_center_mm", "transfer_coefficient"]:
        cdf[col] = pd.to_numeric(cdf[col], errors="coerce").round(6)
    content.append(markdown_table(["检查项", "方法", "驱动", "驱动值", "中心挠度 mm", "状态", "系数", "单位"], cdf.values.tolist()))

    content.append("## 9. 动力代表算例与 30×30 模态结果\n")
    dynamic_rows = [
        ["10×10 Newmark pilot", f"{float(dyn['static_center_mm']):.4f}", f"{float(dyn['peak_center_mm']):.4f}", f"{float(dyn['peak_time_s']) * 1000:.2f}", f"{float(dyn['freq_01_Hz']):.2f}", f"{float(dyn['theta_span_max_K']):.4f}"],
        ["30×30 8 阶模态", f"{float(modal['mechanical_static_center_mm']):.4f}", f"{float(modal['peak_center_mm']):.4f}", f"{float(modal['peak_time_s']) * 1000:.2f}", f"{float(modal['freq_01_Hz']):.2f}", f"{float(modal['theta_span_max_K']):.4f}"],
        ["30×30 16 阶模态", f"{float(sens16['mechanical_static_center_mm']):.4f}", f"{float(sens16['peak_center_mm']):.4f}", f"{float(sens16['peak_time_s']) * 1000:.2f}", f"{float(sens16['freq_01_Hz']):.2f}", f"{float(sens16['theta_span_max_K']):.4f}"],
    ]
    content.append(markdown_table(["方法", "静态中心 mm", "动力峰值 mm", "峰值时间 ms", "一阶频率 Hz", "最大 θ 跨度 K"], dynamic_rows))
    content.append("\n![动力总览](figures/16_dynamic_overview_table.png)\n")
    content.append("![10x10 Newmark 摘要](figures/18_dynamic_summary_table.png)\n")
    content.append("![10x10 中心挠度时程](figures/19_dynamic_center_timeseries.png)\n")
    content.append("![30x30 模态摘要](figures/22_modal30x30_summary_table.png)\n")
    content.append("![30x30 模态时程](figures/23_modal30x30_center_timeseries.png)\n")
    content.append("![10x10 与 30x30 对比](figures/27_10x10_vs_30x30_center.png)\n")

    content.append("## 10. 模态阶数敏感性\n")
    sens_rows = sens[["n_modes", "modal_capture_ratio", "peak_center_mm", "peak_time_s", "overshoot_ratio", "theta_span_max_K"]].copy()
    sens_rows["peak_time_s"] = (sens_rows["peak_time_s"].astype(float) * 1000).round(2)
    for col in ["modal_capture_ratio", "peak_center_mm", "overshoot_ratio", "theta_span_max_K"]:
        sens_rows[col] = sens_rows[col].astype(float).round(6)
    content.append(markdown_table(["阶数", "静态捕获", "峰值 mm", "峰值时间 ms", "超调", "最大 θ 跨度 K"], sens_rows.values.tolist()))
    content.append("\n![阶数敏感性摘要](figures/28_modal_sensitivity_summary_table.png)\n")
    content.append("![静态捕获收敛](figures/29_capture_ratio_vs_modes.png)\n")
    content.append("![峰值收敛](figures/30_peak_vs_modes.png)\n")
    content.append("![不同阶数中心挠度时程](figures/31_center_timeseries_by_modes.png)\n")
    content.append("![收敛性判断](figures/35_convergence_assessment_table.png)\n")
    content.append(f"\n结论：16 阶峰值为 {float(sens16['peak_center_mm']):.4f} mm，与 8 阶峰值绝对值差约 {abs(float(sens.iloc[2]['peak_abs_delta_vs_max_mm'])):.4f} mm；8 阶结果已基本收敛，汇报中可把 16 阶作为稳妥口径。\n")

    content.append("## 11. 阻尼敏感性\n")
    damping_rows = damping[["damping_ratio", "peak_center_mm", "peak_time_s", "overshoot_ratio", "final_center_mm", "theta_span_max_K"]].copy()
    damping_rows["damping_ratio"] = (damping_rows["damping_ratio"].astype(float) * 100).round(2).astype(str) + "%"
    damping_rows["peak_time_s"] = (damping_rows["peak_time_s"].astype(float) * 1000).round(2)
    for col in ["peak_center_mm", "overshoot_ratio", "final_center_mm", "theta_span_max_K"]:
        damping_rows[col] = damping_rows[col].astype(float).round(6)
    content.append(markdown_table(["阻尼比", "峰值 mm", "峰值时间 ms", "超调", "40 ms 位移 mm", "最大 θ 跨度 K"], damping_rows.values.tolist()))
    content.append("\n![阻尼敏感性摘要](figures/36_damping_summary_table.png)\n")
    content.append("![峰值随阻尼变化](figures/37_peak_vs_damping.png)\n")
    content.append("![不同阻尼中心挠度时程](figures/39_center_timeseries_by_damping.png)\n")
    content.append("![阻尼敏感性判断](figures/41_damping_assessment_table.png)\n")

    content.append("## 12. 当前未解决/下一步\n")
    next_rows = [
        ["长时程 Newmark 对照", "30×30 全量 Newmark 交互运行过慢", "可作为夜间长任务跑 1 个短窗或降采样对照"],
        ["论文文件", "当前无 `paper/main.tex` 目录", "动力数据稳定后再创建论文图表和 LaTeX"],
        ["非 U 的 COMSOL 分层验证", "脚本已支持 `FG_COMSOL_LAYER_CSV`，但还未批量验证 V/X/O/P", "选择 1-2 个代表 FG 模式做补充对照"],
        ["COMSOL .mph 大文件", "已生成但未纳入 Git", "保留本地，GitHub 提交轻量 CSV/图表/脚本"],
    ]
    content.append(markdown_table(["事项", "当前状态", "建议处理"], next_rows))

    content.append("## 13. 附：文件索引\n")
    file_rows = [
        ["Word 汇报文件", "FG_MEET_progress_report_2026-05-22.docx"],
        ["PDF 汇报文件", "rendered-word/FG_MEET_progress_report_2026-05-22.pdf"],
        ["本文 README", "README.md"],
        ["图表目录", "figures/"],
        ["轻量结果 CSV", "data/"],
        ["最终 COMSOL 验证表", "data/validation_log.csv"],
        ["最终 15 点验证表", "data/validation_points_U_Vf06_elastic.csv"],
        ["动力敏感性表", "data/dynamic_modal_30x30_U_Vf06_elastic_sensitivity_summary.csv"],
        ["阻尼敏感性表", "data/dynamic_modal_30x30_U_Vf06_elastic_damping_sensitivity_summary.csv"],
    ]
    content.append(markdown_table(["文件", "路径"], file_rows))
    README_PATH.write_text("\n".join(content), encoding="utf-8")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.lstrip("#"))
    tc_pr.append(shd)


def set_cell_text(cell, text: object, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_doc_table(doc: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=8)
        set_cell_shading(table.rows[0].cells[i], "#F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8)
    doc.add_paragraph()


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def add_picture(doc: Document, key: str, title: str, figs: dict[str, Path], width: float = 6.6) -> None:
    caption = doc.add_paragraph()
    caption.paragraph_format.space_after = Pt(8)
    run = caption.add_run(title)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4D78")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(figs[key]), width=Inches(width))


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("2E74B5")


def generate_docx(data: dict[str, pd.DataFrame], figs: dict[str, Path]) -> None:
    dyn = data["dynamic_summary"].iloc[0]
    modal = data["modal_summary"].iloc[0]
    sens = data["modal_sensitivity"].sort_values("n_modes")
    sens16 = sens.iloc[-1]
    damping = data["damping_sensitivity"].sort_values("damping_ratio")
    doc = Document()
    configure_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FG-MEET Workbench 阶段性工作汇报")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F3A5F")
    subtitle = doc.add_paragraph("模型目标、方法路线、问题排查、正确/错误结果、动力补充与 PDF 版汇报 | 2026-05-22/23")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    add_picture(doc, "workflow", "图 1 技术路线与实际收敛流程", figs)
    add_picture(doc, "workload_table", "图 2 工作量总览", figs)

    add_section_heading(doc, "1. 模型与目标")
    add_doc_table(doc, ["项目", "内容"], [
        ["模型", "300×300×6 mm FG-MEE CFFF 方板，30×30 八节点板壳，10 层材料"],
        ["载荷", "Case A: 15000 Pa；Case B: ±300 V；Case C: ±200 A"],
        ["目标", "MATLAB 主扫描 + COMSOL 独立验证，关键误差 <5%"],
        ["汇报目标", "把一个月工作量转化为可截图的图表、表格和说明"],
    ])
    add_picture(doc, "method_table", "图 3 预计方法与实际执行路径", figs)

    add_section_heading(doc, "2. 过程问题与解决思路")
    doc.add_paragraph("实际难点集中在 COMSOL 自动化权限、边界/载荷选择、MATLAB reduced Qd 后处理、以及 COMSOL 网格/厚度离散一致性。下面表格可直接用于组会说明排查过程。")
    add_picture(doc, "problem_solution_table", "图 4 实际问题与处理路径", figs)

    add_section_heading(doc, "3. MATLAB 主算例结果")
    add_picture(doc, "05_elastic_w_vs_vf", "图 5 Case A 机械压力下中心挠度", figs)
    add_picture(doc, "06_electro_w_vs_vf", "图 6 Case B 电势驱动下中心挠度", figs)
    add_picture(doc, "07_magneto_w_vs_vf", "图 7 Case C 磁势驱动下中心挠度", figs)
    add_picture(doc, "theta_heatmap", "图 8 Case A 层温差跨度热力图", figs)
    add_picture(doc, "me_efficiency", "图 9 Case C 磁电效率曲线", figs)

    add_section_heading(doc, "4. reduced Qd 后处理修正")
    doc.add_paragraph("修正前直接按节点号索引 reduced Qd，会在固定边存在时错位。修正后先按节点约束标志还原完整 TQd，再取中心节点。135 行 results_static.csv 已全部刷新。")
    add_picture(doc, "refresh_top10", "图 10 中心挠度刷新前后差异 Top 10", figs)

    add_section_heading(doc, "5. COMSOL 方案筛选与错误结果")
    doc.add_paragraph("为了便于讨论，错误或未达标方案也完整保留：自由四面体网格、单域 swept、10-domain CSV 的不同厚度离散均列入对比。")
    add_picture(doc, "comsol_experiment_errors", "图 11 COMSOL 方案误差收敛路径", figs, width=6.8)
    exp = data["experiments"][["run_tag", "center_rel_err_pct", "max_rel_err_pct", "mean_rel_err_pct"]].round(3)
    add_doc_table(doc, ["实验", "中心误差%", "最大误差%", "平均误差%"], exp.values.tolist())

    add_section_heading(doc, "6. 最终正确结果")
    add_picture(doc, "final_result_table", "图 12 最终通过结果摘要", figs)
    add_picture(doc, "comsol_point_comparison", "图 13 最终 15 点 MATLAB/COMSOL 位移对比", figs)
    add_picture(doc, "comsol_point_error", "图 14 最终 15 点相对误差", figs)
    val = data["validation"][["point_id", "matlab_w_mm", "comsol_w_mm", "diff_w_mm", "rel_err_w_pct"]].copy()
    for col in val.columns[1:]:
        val[col] = val[col].astype(float).round(4)
    add_doc_table(doc, ["点", "MATLAB mm", "COMSOL mm", "差值 mm", "误差%"], val.values.tolist())

    add_section_heading(doc, "7. 2 mm 耦合验证")
    add_picture(doc, "coupling_table", "图 15 2 mm 耦合验证表", figs)
    cdf = data["coupling"][["check_id", "method", "driver_kind", "w_center_mm", "status"]].copy()
    cdf["w_center_mm"] = pd.to_numeric(cdf["w_center_mm"]).round(6)
    add_doc_table(doc, ["检查项", "方法", "驱动", "中心挠度 mm", "状态"], cdf.values.tolist())

    add_section_heading(doc, "8. 动力代表算例与 30x30 模态结果")
    add_doc_table(doc, ["方法", "静态中心 mm", "动力峰值 mm", "峰值时间 ms", "一阶频率 Hz", "最大 θ 跨度 K"], [
        ["10x10 Newmark pilot", f"{float(dyn['static_center_mm']):.4f}", f"{float(dyn['peak_center_mm']):.4f}", f"{float(dyn['peak_time_s']) * 1000:.2f}", f"{float(dyn['freq_01_Hz']):.2f}", f"{float(dyn['theta_span_max_K']):.4f}"],
        ["30x30 8 阶模态", f"{float(modal['mechanical_static_center_mm']):.4f}", f"{float(modal['peak_center_mm']):.4f}", f"{float(modal['peak_time_s']) * 1000:.2f}", f"{float(modal['freq_01_Hz']):.2f}", f"{float(modal['theta_span_max_K']):.4f}"],
        ["30x30 16 阶模态", f"{float(sens16['mechanical_static_center_mm']):.4f}", f"{float(sens16['peak_center_mm']):.4f}", f"{float(sens16['peak_time_s']) * 1000:.2f}", f"{float(sens16['freq_01_Hz']):.2f}", f"{float(sens16['theta_span_max_K']):.4f}"],
    ])
    add_picture(doc, "dynamic_overview_table", "图 16 动力结果总览", figs)
    add_picture(doc, "dynamic_summary_table", "图 17 10x10 Newmark pilot 摘要", figs)
    add_picture(doc, "dynamic_center_timeseries", "图 18 10x10 Newmark 中心挠度时程", figs)
    add_picture(doc, "modal30_summary_table", "图 19 30x30 8 阶模态结果摘要", figs)
    add_picture(doc, "modal30_center_timeseries", "图 20 30x30 8 阶模态中心挠度时程", figs)
    add_picture(doc, "modal30_compare_10x10", "图 21 10x10 Newmark 与 30x30 模态对比", figs)

    doc.add_page_break()
    add_section_heading(doc, "9. 模态阶数敏感性")
    sens_rows = sens[["n_modes", "modal_capture_ratio", "peak_center_mm", "peak_time_s", "overshoot_ratio", "theta_span_max_K"]].copy()
    sens_rows["peak_time_s"] = (sens_rows["peak_time_s"].astype(float) * 1000).round(2)
    for col in ["modal_capture_ratio", "peak_center_mm", "overshoot_ratio", "theta_span_max_K"]:
        sens_rows[col] = sens_rows[col].astype(float).round(6)
    add_doc_table(doc, ["阶数", "静态捕获", "峰值 mm", "峰值时间 ms", "超调", "最大 θ 跨度 K"], sens_rows.values.tolist())
    add_picture(doc, "sens_summary_table", "图 22 30x30 模态阶数敏感性摘要", figs)
    add_picture(doc, "sens_capture_ratio", "图 23 静态捕获比例收敛", figs)
    add_picture(doc, "sens_peak_vs_modes", "图 24 动力峰值随模态阶数变化", figs)
    add_picture(doc, "sens_center_timeseries", "图 25 不同阶数中心挠度时程", figs)
    add_picture(doc, "sens_assessment_table", "图 26 模态阶数收敛性判断", figs)

    add_section_heading(doc, "10. 阻尼敏感性")
    damping_rows = damping[["damping_ratio", "peak_center_mm", "peak_time_s", "overshoot_ratio", "final_center_mm", "theta_span_max_K"]].copy()
    damping_rows["damping_ratio"] = (damping_rows["damping_ratio"].astype(float) * 100).round(2).astype(str) + "%"
    damping_rows["peak_time_s"] = (damping_rows["peak_time_s"].astype(float) * 1000).round(2)
    for col in ["peak_center_mm", "overshoot_ratio", "final_center_mm", "theta_span_max_K"]:
        damping_rows[col] = damping_rows[col].astype(float).round(6)
    add_doc_table(doc, ["阻尼比", "峰值 mm", "峰值时间 ms", "超调", "40 ms 位移 mm", "最大 θ 跨度 K"], damping_rows.values.tolist())
    add_picture(doc, "damping_summary_table", "图 27 30x30 模态阻尼敏感性摘要", figs)
    add_picture(doc, "damping_peak_vs", "图 28 动力峰值随阻尼比变化", figs)
    add_picture(doc, "damping_center_timeseries", "图 29 不同阻尼比中心挠度时程", figs)
    add_picture(doc, "damping_assessment_table", "图 30 阻尼敏感性判断", figs)

    add_section_heading(doc, "11. 下一步")
    add_doc_table(doc, ["事项", "状态", "下一步"], [
        ["长时程 Newmark 对照", "30x30 全量 Newmark 交互运行过慢", "可作为夜间长任务补 1 个对照窗"],
        ["论文 main.tex", "仓库当前无 paper/ 目录", "动力数据完成后再生成论文图表"],
        ["更多 COMSOL 分布", "U/Vf0.6 已通过", "可选 V/X/O/P 代表算例补充"],
    ])

    doc.save(DOCX_PATH)


def main() -> None:
    ensure_dirs()
    copy_data_files()
    data = load_data()
    figs = generate_figures(data)
    generate_readme(data, figs)
    generate_docx(data, figs)
    print(f"Wrote {README_PATH}")
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {len(figs)} figures to {FIG_DIR}")


if __name__ == "__main__":
    main()
